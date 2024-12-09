from flask import Flask, render_template, request, send_file
import pandas as pd
from io import BytesIO
from docx import Document

app = Flask(__name__)

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/upload", methods=["POST"])
def upload_file():
    if "file" not in request.files:
        return "No file provided", 400

    file = request.files["file"]

    try:
        # Load the file into a Pandas DataFrame
        if file.filename.endswith(".csv"):
            df = pd.read_csv(file)
        elif file.filename.endswith(".xlsx"):
            df = pd.read_excel(file)
        else:
            return "Unsupported file format. Please upload .csv or .xlsx files.", 400
        
        # Create a DOCX document
        doc = Document()
        doc.add_heading("Data Table", level=1)
        
        # Add table to the DOCX
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'

        # Add column headers
        hdr_cells = table.rows[0].cells
        for i, column_name in enumerate(df.columns):
            hdr_cells[i].text = str(column_name)

        # Add data rows
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)

        # Save to a BytesIO stream
        output = BytesIO()
        doc.save(output)
        output.seek(0)

        return send_file(
            output,
            as_attachment=True,
            download_name="converted.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        return str(e), 500


if __name__ == "__main__":
    app.run(debug=True)
