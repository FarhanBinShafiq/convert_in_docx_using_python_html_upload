from flask import Flask, render_template, request, send_file
import pandas as pd
from io import BytesIO
from docx import Document

app = Flask(__name__)

# Route to render the HTML interface
@app.route("/")
def index():
    return render_template("index.html")

# Route to handle file uploads and conversion
@app.route("/upload", methods=["POST"])
def upload_file():
    if "file" not in request.files:
        return "No file provided", 400

    file = request.files["file"]

    try:
        # Load the file into Pandas DataFrame
        if file.filename.endswith(".csv"):
            df = pd.read_csv(file)
        elif file.filename.endswith(".xlsx"):
            df = pd.read_excel(file)
        else:
            return "Unsupported file format. Please upload .csv or .xlsx files.", 400
        
        # Create a DOCX document
        doc = Document()
        doc.add_heading("Converted Data", level=1)
        for index, row in df.iterrows():
            doc.add_paragraph(", ".join(map(str, row.values)))

        # Save to a BytesIO stream
        output = BytesIO()
        doc.save(output)
        output.seek(0)

        return send_file(
            output,
            as_attachment=True,
            download_name="converted.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        return str(e), 500


if __name__ == "__main__":
    app.run(debug=True)
